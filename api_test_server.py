from flask import Flask, request, jsonify, send_file
import os
import json
import tempfile
from docx import Document
import secrets
import base64

# Generate a random API key if one doesn't exist in environment variables
API_KEY = os.environ.get('API_KEY', 'docgen_api_12345')
print(f"API Key: {API_KEY}")  # Print the API key when the app starts

# Create Flask app for API endpoints
app = Flask(__name__)

def check_api_key():
    """Check if the API key is valid."""
    provided_key = request.headers.get('X-API-Key')
    if not provided_key or provided_key != API_KEY:
        return False
    return True

@app.route('/', methods=['GET'])
def index():
    """Index route to check if the API is running."""
    return jsonify({
        "status": "running",
        "endpoints": [
            {"path": "/api/docx-to-json", "method": "POST", "description": "Convert DOCX to JSON"},
            {"path": "/api/json-to-docx", "method": "POST", "description": "Convert JSON to DOCX"}
        ],
        "note": "API key required in X-API-Key header"
    })

@app.route('/api/docx-to-json', methods=['POST'])
def api_docx_to_json():
    # Check API key
    if not check_api_key():
        return jsonify({"error": "Invalid or missing API key"}), 401
    
    # Determine the request content type
    content_type = request.headers.get('Content-Type', '')
    
    try:
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, "document.docx")
        
        # Handle different request formats
        if 'multipart/form-data' in content_type:
            # Traditional file upload
            if 'file' not in request.files:
                return jsonify({"error": "No file part"}), 400
            
            file = request.files['file']
            if file.filename == '':
                return jsonify({"error": "No selected file"}), 400
            
            if not file.filename.lower().endswith('.docx'):
                return jsonify({"error": "File must be a DOCX document"}), 400
            
            file.save(file_path)
            
        elif 'application/json' in content_type:
            # Base64 encoded file
            if not request.is_json:
                return jsonify({"error": "Invalid JSON"}), 400
                
            json_data = request.json
            if 'base64_content' not in json_data:
                return jsonify({"error": "Missing base64_content field"}), 400
                
            try:
                file_content = base64.b64decode(json_data['base64_content'])
                with open(file_path, 'wb') as f:
                    f.write(file_content)
            except Exception as e:
                return jsonify({"error": f"Invalid base64 content: {str(e)}"}), 400
                
        elif 'application/octet-stream' in content_type:
            # Raw binary data
            file_content = request.data
            if not file_content:
                return jsonify({"error": "Empty binary content"}), 400
                
            with open(file_path, 'wb') as f:
                f.write(file_content)
                
        else:
            return jsonify({"error": "Unsupported content type. Use multipart/form-data, application/json with base64_content, or application/octet-stream"}), 400
        
        # Simple DOCX to JSON conversion for testing
        doc = Document(file_path)
        result = {
            "document": {
                "blocks": []
            }
        }
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                block = {
                    "type": "paragraph",
                    "runs": [
                        {
                            "text": para.text,
                            "bold": False,
                            "italic": False,
                            "underline": False
                        }
                    ]
                }
                result["document"]["blocks"].append(block)
        
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/json-to-docx', methods=['POST'])
def api_json_to_docx():
    # Check API key
    if not check_api_key():
        return jsonify({"error": "Invalid or missing API key"}), 401
        
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 400
    
    try:
        # Get the JSON data
        json_data = request.json
        
        # Create a simple DOCX document
        doc = Document()
        
        # Add content from JSON
        if "document" in json_data and "blocks" in json_data["document"]:
            for block in json_data["document"]["blocks"]:
                if block.get("type") == "paragraph":
                    p = doc.add_paragraph()
                    for run in block.get("runs", []):
                        text = run.get("text", "")
                        r = p.add_run(text)
                        if run.get("bold"):
                            r.bold = True
                        if run.get("italic"):
                            r.italic = True
                        if run.get("underline"):
                            r.underline = True
                elif block.get("type") == "heading":
                    level = block.get("level", 1)
                    text = " ".join([run.get("text", "") for run in block.get("runs", [])])
                    doc.add_heading(text, level=level)
        
        # Save the document
        temp_dir = tempfile.mkdtemp()
        docx_path = os.path.join(temp_dir, "converted.docx")
        doc.save(docx_path)
        
        # Make sure the file exists before sending
        if not os.path.exists(docx_path):
            return jsonify({"error": "Failed to create DOCX file"}), 500
            
        print(f"DOCX file created at: {docx_path}")
        
        # Use send_file with the correct mimetype
        return send_file(
            docx_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name="converted.docx"
        )
    except Exception as e:
        print(f"Error in JSON to DOCX conversion: {str(e)}")
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    # Run the Flask app
    port = int(os.environ.get('PORT', 9999))
    print(f"Starting API server on http://localhost:{port}")
    app.run(host='0.0.0.0', port=port, debug=True)
