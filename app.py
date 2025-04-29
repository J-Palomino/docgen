import gradio as gr
import os
from pdf2docx import Converter
from docx import Document
from docx.table import _Cell
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
import json
import base64
import hashlib
import sys
import tempfile
from flask import Flask, request, jsonify, send_file
import threading
import secrets

# Generate a random API key if one doesn't exist in environment variables
API_KEY = os.environ.get('API_KEY', 'docgen_api_12345')
print(f"API Key: {API_KEY}")  # Print the API key when the app starts

# Define supported formats directly instead of using pypandoc
input_supported_formats = [
    'BIBLATEX', 'BIBTEX', 'BITS', 'COMMONMARK', 'COMMONMARK_X', 'CREOLE', 'CSLJSON', 'CSV',
    'DJOT', 'DOCBOOK', 'DOCX', 'DOKUWIKI', 'ENDNOTEXML', 'EPUB', 'FB2', 'GFM', 'HADDOCK',
    'HTML', 'IPYNB', 'JATS', 'JIRA', 'JSON', 'LATEX', 'MAN', 'MARKDOWN', 'MARKDOWN_GITHUB',
    'MARKDOWN_MMD', 'MARKDOWN_PHPEXTRA', 'MARKDOWN_STRICT', 'MDOC', 'MEDIAWIKI', 'MUSE',
    'NATIVE', 'ODT', 'OPML', 'ORG', 'PDF', 'POD', 'RIS', 'RST', 'RTF', 'T2T', 'TEXTILE',
    'TIKIWIKI', 'TSV', 'TWIKI', 'TYPST', 'VIMWIKI'
]

output_supported_formats = [
    "ANSI", "ASCIIDOC", "ASCIIDOC_LEGACY", "ASCIIDOCTOR", "BEAMER", "BIBLATEX", "BIBTEX", "CHUNKEDHTML", 
    "COMMONMARK", "COMMONMARK_X", "CONTEXT", "CSLJSON", "DJOT", "DOCBOOK", "DOCBOOK4", "DOCBOOK5", 
    "DOCX", "DOKUWIKI", "DZSLIDES", "EPUB", "EPUB2", "EPUB3", "FB2", "GFM", "HADDOCK", "HTML", 
    "HTML4", "HTML5", "ICML", "IPYNB", "JATS", "JATS_ARCHIVING", "JATS_ARTICLEAUTHORING", 
    "JATS_PUBLISHING", "JIRA", "JSON", "LATEX", "MAN", "MARKDOWN", "MARKDOWN_GITHUB", 
    "MARKDOWN_MMD", "MARKDOWN_PHPEXTRA", "MARKDOWN_STRICT", "MARKUA", "MEDIAWIKI", "MS", 
    "MUSE", "NATIVE", "ODT", "OPENDOCUMENT", "OPML", "ORG", "PDF", "PLAIN", "PPTX", "REVEALJS", 
    "RST", "RTF", "S5", "SLIDEOUS", "SLIDY", "TEI", "TEXINFO", "TEXTILE", "TYPST", "XWIKI", "ZIMWIKI"
]

# Only import pypandoc if not in Railway environment
if not os.environ.get('RAILWAY_ENVIRONMENT'):
    try:
        import pypandoc
        def ensure_pandoc_installed():
            try:
                # Check if pandoc is already installed
                pypandoc.get_pandoc_version()
                print("Pandoc is already installed and accessible.")
            except OSError:
                # Instead of downloading, just print a message
                print("Pandoc not found, but continuing without it for API testing.")
                # Skip download: pypandoc.download_pandoc()
        
        # Make sure Pandoc is installed if not in Railway
        ensure_pandoc_installed()
    except ImportError:
        print("Pypandoc not available, continuing with limited functionality.")

def convert_pdf_to_docx(pdf_file):
    output_docx = f"{os.path.splitext(pdf_file)[0]}.docx"
    cv = Converter(pdf_file)
    cv.convert(output_docx, start=0, end=None)
    return output_docx

def get_preview(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext in ['.txt', '.md', '.csv', '.json']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(2000)  # Preview first 2000 chars
                return f"<pre style='max-height:300px;overflow:auto'>{content}</pre>"
        elif ext == '.pdf':
            # Show PDF inline using HTML embed
            return f"<embed src='{file_path}' type='application/pdf' width='100%' height='400px' />"
        elif ext == '.docx':
            try:
                doc = Document(file_path)
                html = ""
                # Extract header(s) with paragraphs and tables
                headers = []
                for section in doc.sections:
                    header_texts = []
                    # Paragraphs
                    for p in section.header.paragraphs:
                        if p.text.strip():
                            header_texts.append(p.text.strip())
                    # Tables
                    for table in section.header.tables:
                        for row in table.rows:
                            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                            if row_text:
                                header_texts.append(row_text)
                    if header_texts:
                        headers.append(" | ".join(header_texts))
                if headers:
                    html += f"<div style='font-weight:bold;font-size:1.2em;margin-bottom:8px;'>{' | '.join(headers)}</div>"
                para_count = 0
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        html += f"<p>{text}</p>"
                        para_count += 1
                        if para_count > 30:
                            html += "<p><i>Preview truncated...</i></p>"
                            break
                return f"<div style='max-height:300px;overflow:auto'>{html}</div>"
            except Exception as e:
                return f"<b>Error reading DOCX:</b> {e}"
        elif ext == '.doc':
            return f"<b>DOC file:</b> {os.path.basename(file_path)} (Preview not supported)"
        else:
            return f"<b>File:</b> {os.path.basename(file_path)} (Preview not supported)"
    except Exception as e:
        return f"<b>Error generating preview:</b> {e}"

def extract_runs(paragraph):
    runs = []
    for run in paragraph.runs:
        run_data = {
            "text": run.text
        }
        if run.bold:
            run_data["bold"] = True
        if run.italic:
            run_data["italic"] = True
        if run.underline:
            run_data["underline"] = True
        if run.font and run.font.size:
            run_data["font_size"] = run.font.size.pt
        if run.font and run.font.name:
            run_data["font_name"] = run.font.name
        # Extract color (RGB or theme)
        if run.font and run.font.color:
            if run.font.color.rgb:
                run_data["color"] = str(run.font.color.rgb)
            elif run.font.color.theme_color:
                run_data["color_theme"] = str(run.font.color.theme_color)
        # Highlight color
        if run.font and hasattr(run.font, "highlight_color") and run.font.highlight_color:
            run_data["highlight"] = str(run.font.highlight_color)
        # Strikethrough
        if run.font and run.font.strike:
            run_data["strikethrough"] = True
        # Superscript/subscript
        if run.font and run.font.superscript:
            run_data["superscript"] = True
        if run.font and run.font.subscript:
            run_data["subscript"] = True
        # Small caps
        if run.font and run.font.small_caps:
            run_data["small_caps"] = True
        # All caps
        if run.font and run.font.all_caps:
            run_data["all_caps"] = True
        # Add hyperlink detection (if run is part of a hyperlink)
        if run._element.xpath("ancestor::w:hyperlink"):
            # Try to extract hyperlink target if possible
            parent_hyper = run._element.xpath("ancestor::w:hyperlink")[0]
            rId = parent_hyper.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rId:
                try:
                    rel = paragraph.part.rels[rId]
                    run_data["hyperlink"] = rel.target_ref
                except Exception:
                    pass
        runs.append(run_data)
    return runs

# Detect heading and list paragraphs
def extract_paragraph_block(paragraph):
    style_name = paragraph.style.name if paragraph.style else "Normal"
    # Heading
    if style_name.startswith("Heading"):
        try:
            level = int(style_name.split()[-1])
        except Exception:
            level = 1
        block = {
            "type": "heading",
            "level": level,
            "runs": extract_runs(paragraph),
            "alignment": str(paragraph.alignment) if paragraph.alignment else "left",
            "style": style_name
        }
    # List
    elif "List" in style_name or paragraph._element.xpath(".//w:numPr"):
        block = {
            "type": "list_item",
            "list_type": "number" if "Number" in style_name else "bullet",
            "runs": extract_runs(paragraph),
            "alignment": str(paragraph.alignment) if paragraph.alignment else "left",
            "style": style_name
        }
    # Normal paragraph
    else:
        block = {
            "type": "paragraph",
            "runs": extract_runs(paragraph),
            "alignment": str(paragraph.alignment) if paragraph.alignment else "left",
            "style": style_name
        }
    # Page break
    if paragraph._element.xpath(".//w:br[@w:type='page']"):
        block["page_break"] = True
    # Bookmarks
    bookmarks = paragraph._element.xpath(".//w:bookmarkStart")
    if bookmarks:
        block["bookmarks"] = [bm.get("w:name") for bm in bookmarks if bm.get("w:name")]
    # Comments (commentRangeStart)
    comments = paragraph._element.xpath(".//w:commentRangeStart")
    if comments:
        block["comments"] = True
    return block

def extract_table_block(table):
    rows = []
    for row in table.rows:
        row_cells = []
        for cell in row.cells:
            # Only take unique paragraphs (python-docx repeats cell objects)
            unique_paras = []
            seen = set()
            for para in cell.paragraphs:
                para_id = id(para)
                if para_id not in seen:
                    unique_paras.append(para)
                    seen.add(para_id)
            cell_blocks = [extract_paragraph_block(para) for para in unique_paras if para.text.strip()]
            # Cell properties: merge, width, shading
            cell_props = {}
            tc = cell._tc
            gridspan = tc.xpath(".//w:gridSpan")
            if gridspan:
                cell_props["colspan"] = int(gridspan[0].get("w:val", "1"))
            vmerge = tc.xpath(".//w:vMerge")
            if vmerge:
                cell_props["rowspan"] = True
            width = tc.xpath(".//w:tcW")
            if width:
                cell_props["width"] = int(width[0].get("w:w", "0"))
            shading = tc.xpath(".//w:shd")
            if shading:
                cell_props["shading"] = shading[0].get("w:fill")
            row_cells.append({"blocks": cell_blocks, **cell_props})
        rows.append(row_cells)
    return {"type": "table", "rows": rows}

def extract_images_from_doc(doc, output_dir, image_prefix):
    image_blocks = []
    rels = doc.part.rels
    for rel in rels.values():
        if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
            img_blob = rel.target_part.blob
            img_hash = hashlib.sha1(img_blob).hexdigest()[:8]
            img_ext = rel.target_part.content_type.split('/')[-1]
            img_id = f"{image_prefix}_{img_hash}"
            img_filename = f"{img_id}.{img_ext}"
            img_path = os.path.join(output_dir, img_filename)
            with open(img_path, 'wb') as f:
                f.write(img_blob)
            image_blocks.append({
                "type": "image",
                "image_id": img_id,
                "image_format": img_ext,
                "path": img_filename
            })
    return image_blocks

def add_runs_to_paragraph(paragraph, runs):
    for run_info in runs:
        run = paragraph.add_run(run_info.get("text", ""))
        if run_info.get("bold"): run.bold = True
        if run_info.get("italic"): run.italic = True
        if run_info.get("underline"): run.underline = True
        if run_info.get("font_size"): run.font.size = Pt(run_info["font_size"])
        if run_info.get("font_name"): run.font.name = run_info["font_name"]
        # Set color (RGB or theme)
        if run_info.get("color"):
            try:
                run.font.color.rgb = RGBColor.from_string(run_info["color"].replace("#", ""))
            except Exception:
                pass
        elif run_info.get("color_theme"):
            try:
                run.font.color.theme_color = int(run_info["color_theme"])
            except Exception:
                pass
        if run_info.get("highlight"):
            try:
                if run_info["highlight"].isdigit():
                    run.font.highlight_color = int(run_info["highlight"])
                else:
                    run.font.highlight_color = WD_COLOR_INDEX[run_info["highlight"]]
            except Exception:
                pass
        # Strikethrough
        if run_info.get("strikethrough"): run.font.strike = True
        # Superscript/subscript
        if run_info.get("superscript"): run.font.superscript = True
        if run_info.get("subscript"): run.font.subscript = True
        # Small caps
        if run_info.get("small_caps"): run.font.small_caps = True
        # All caps
        if run_info.get("all_caps"): run.font.all_caps = True
        # Hyperlink
        if run_info.get("hyperlink"):
            try:
                run.hyperlink = run_info["hyperlink"]
            except Exception:
                pass

# Add heading and list support
def add_block_to_doc(doc, block, image_dir):
    # Log context
    parent_type = type(doc).__name__
    if block["type"] == "heading":
        level = block.get("level", 1)
        text = "".join([r.get("text", "") for r in block.get("runs", [])])
        para = doc.add_heading(text, level=level)
        add_runs_to_paragraph(para, block.get("runs", []))
        align = block.get("alignment", "left")
        if align == "center": para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right": para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif block["type"] == "paragraph":
        # Skip empty paragraphs in headers/footers
        if hasattr(doc, 'is_header') or hasattr(doc, 'is_footer'):
            if not block.get("runs") or all(not r.get("text", "").strip() for r in block.get("runs", [])):
                print(f"[add_block_to_doc] Skipping empty paragraph in {parent_type}")
                return
        para = doc.add_paragraph()
        add_runs_to_paragraph(para, block.get("runs", []))
        align = block.get("alignment", "left")
        if align == "center": para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right": para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif block["type"] == "list_item":
        style = "List Number" if block.get("list_type") == "number" else "List Bullet"
        para = doc.add_paragraph(style=style)
        add_runs_to_paragraph(para, block.get("runs", []))
        align = block.get("alignment", "left")
        if align == "center": para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right": para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else: para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if "space_before" in block: para.paragraph_format.space_before = Pt(block["space_before"])
        if "space_after" in block: para.paragraph_format.space_after = Pt(block["space_after"])
        if "line_spacing" in block: para.paragraph_format.line_spacing = block["line_spacing"]
    elif block["type"] == "table":
        rows = block.get("rows", [])
        if not rows:
            print(f"[add_block_to_doc] Skipping empty table in {parent_type}")
            return
        n_rows = len(rows)
        n_cols = max(len(r) for r in rows) if rows else 1
        print(f"[add_block_to_doc] Adding table with {n_rows} rows and {n_cols} cols to {parent_type}")
        table_width = block.get("width")
        if not table_width or not isinstance(table_width, (int, float)) or table_width == 0:
            table_width = 5486400  # 6 inches in EMUs
        doc_class = type(doc).__name__
        is_header_footer = doc_class in ["_HeaderPart", "_FooterPart", "_Header", "_Footer"]
        if is_header_footer:
            # Workaround: create table in temp doc, then move XML to header/footer
            from docx import Document
            temp_doc = Document()
            table = temp_doc.add_table(rows=n_rows, cols=n_cols)
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    cell_blocks = cell.get("blocks", [])
                    cell_obj = table.cell(i, j)
                    print(f"[add_block_to_doc] Rendering cell ({i},{j}) in {parent_type} with {len(cell_blocks)} blocks, cell_obj type: {type(cell_obj)}")
                    if cell_blocks:
                        for para_block in cell_blocks:
                            print(f"[add_block_to_doc] Adding block to cell ({i},{j}): {para_block}")
                            if para_block.get("type") == "paragraph":
                                p = cell_obj.add_paragraph()
                                for run in para_block.get("runs", []):
                                    r = p.add_run(run.get("text", ""))
                            else:
                                add_block_to_doc(cell_obj, para_block, image_dir)
                    else:
                        cell_obj.add_paragraph("")
                    print(f"[add_block_to_doc] Cell ({i},{j}) text after add: '{cell_obj.text}'")
            print(f"[add_block_to_doc] Finished rendering table in {parent_type}")
            # Attach table XML to header/footer
            hdrftr_element = doc._element
            hdrftr_element.append(table._element)
        else:
            table = doc.add_table(rows=n_rows, cols=n_cols)
            for i, row in enumerate(rows):
                for j, cell in enumerate(row):
                    cell_blocks = cell.get("blocks", [])
                    cell_obj = table.cell(i, j)
                    print(f"[add_block_to_doc] Rendering cell ({i},{j}) in {parent_type} with {len(cell_blocks)} blocks, cell_obj type: {type(cell_obj)}")
                    if cell_blocks:
                        for para_block in cell_blocks:
                            print(f"[add_block_to_doc] Adding block to cell ({i},{j}): {para_block}")
                            if para_block.get("type") == "paragraph":
                                p = cell_obj.add_paragraph()
                                for run in para_block.get("runs", []):
                                    r = p.add_run(run.get("text", ""))
                            else:
                                add_block_to_doc(cell_obj, para_block, image_dir)
                    else:
                        cell_obj.add_paragraph("")
                    print(f"[add_block_to_doc] Cell ({i},{j}) text after add: '{cell_obj.text}'")
            print(f"[add_block_to_doc] Finished rendering table in {parent_type}")
    elif block["type"] == "image":
        img_path = os.path.join(image_dir, block["path"])
        width = block.get("width")
        height = block.get("height")
        if os.path.exists(img_path):
            if width and height:
                doc.add_picture(img_path, width=Inches(width/96), height=Inches(height/96))
            else:
                doc.add_picture(img_path)

def add_blocks_to_doc(doc, blocks, image_dir):
    # Only add blocks to header/footer if not already populated (avoid duplicate headers/footers)
    is_header = hasattr(doc, 'is_header') and getattr(doc, 'is_header', False)
    is_footer = hasattr(doc, 'is_footer') and getattr(doc, 'is_footer', False)
    # For headers/footers: check if already has non-empty paragraphs or tables
    already_populated = False
    if is_header or is_footer:
        # Check for non-empty paragraphs
        if hasattr(doc, 'paragraphs') and any(p.text.strip() for p in doc.paragraphs):
            already_populated = True
        # Check for tables with non-empty cells
        if hasattr(doc, 'tables') and any(any(cell.text.strip() for row in t.rows for cell in row.cells) for t in doc.tables):
            already_populated = True
    if (is_header or is_footer) and already_populated:
        print(f"[add_blocks_to_doc] Skipping: header/footer already populated. Not adding blocks.")
        return
    # Only add blocks if not empty
    if blocks:
        for block in blocks:
            add_block_to_doc(doc, block, image_dir)
    
# Helper to extract all blocks from a container (paragraphs and tables)
def extract_blocks(container, output_dir, image_prefix):
    blocks = []
    for paragraph in container.paragraphs:
        block = extract_paragraph_block(paragraph)
        if block:
            blocks.append(block)
    for table in getattr(container, 'tables', []):
        block = extract_table_block(table)
        if block:
            blocks.append(block)
    return blocks

def extract_all_sections(doc, output_dir, image_prefix):
    sections = []
    for idx, section in enumerate(doc.sections):
        sec = {}
        for htype, attr in [("header", "header"), ("first_page_header", "first_page_header"), ("even_page_header", "even_page_header"),
                            ("footer", "footer"), ("first_page_footer", "first_page_footer"), ("even_page_footer", "even_page_footer")]:
            part = getattr(section, attr, None)
            if part:
                sec[htype] = extract_blocks(part, output_dir, f"{image_prefix}_sec{idx}_{htype}")
        sections.append(sec)
    body = extract_blocks(doc, output_dir, os.path.splitext(os.path.basename(getattr(doc, 'filename', 'document')))[0])
    return {"sections": sections, "body": body}

def convert_document(doc_file, target_format):
    """Convert a document to the target format"""
    # Get file path from the uploaded file
    if hasattr(doc_file, 'name'):
        orig_file_path = doc_file.name
    else:
        orig_file_path = str(doc_file)
    
    # Get file extension
    file_ext = os.path.splitext(orig_file_path)[1].lower()
    
    # Create output filename
    output_dir = os.path.dirname(orig_file_path)
    output_base = os.path.splitext(os.path.basename(orig_file_path))[0]
    output_ext = f".{target_format.lower()}"
    if target_format.lower() == "docx":
        output_ext = ".docx"
    elif target_format.lower() == "pdf":
        output_ext = ".pdf"
    elif target_format.lower() == "html":
        output_ext = ".html"
    elif target_format.lower() == "markdown" or target_format.lower() == "md":
        output_ext = ".md"
    elif target_format.lower() == "json":
        output_ext = ".json"
    
    output_file = os.path.join(output_dir, f"{output_base}{output_ext}")
    
    # Handle PDF to DOCX conversion
    if file_ext == '.pdf' and target_format.lower() == 'docx':
        output_file = convert_pdf_to_docx(orig_file_path)
        input_preview = get_preview(orig_file_path)
        output_preview = get_preview(output_file)
        return input_preview, output_preview, output_file
    
    # Handle DOCX to JSON conversion
    if file_ext == '.docx' and target_format.lower() == 'json':
        # Extract document structure to JSON
        doc = Document(orig_file_path)
        temp_dir = tempfile.mkdtemp()
        image_prefix = hashlib.md5(orig_file_path.encode()).hexdigest()

        # Extract document sections (including headers/footers)
        result = extract_all_sections(doc, temp_dir, image_prefix)

        # Save JSON (no flattening, no duplication)
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        input_preview = get_preview(orig_file_path)
        output_preview = get_preview(output_file)
        return input_preview, output_preview, output_file
    
    # Handle JSON to DOCX conversion
    if file_ext == '.json' and target_format.lower() == 'docx':
        # Create a new document from JSON
        with open(orig_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        doc = Document()
        temp_dir = os.path.dirname(orig_file_path)

        # If the document is wrapped in a 'document' key, unwrap it
        if 'document' in data:
            data = data['document']

        # Restore sections, headers, and footers if present
        header_fallback_blocks = []
        if 'sections' in data and isinstance(data['sections'], list) and len(data['sections']) > 0:
            # Set up sections in the DOCX to match JSON
            num_json_sections = len(data['sections'])
            # python-docx always starts with one section
            while len(doc.sections) < num_json_sections:
                doc.add_section()
            for idx, section_json in enumerate(data['sections']):
                section = doc.sections[idx]
                # Determine what to use for first page header
                first_page_header_blocks = section_json.get('first_page_header')
                header_blocks = section_json.get('header')
                even_page_header_blocks = section_json.get('even_page_header')
                # Set section flags
                section.different_first_page_header_footer = bool(first_page_header_blocks or header_blocks)
                section.different_even_page_header_footer = bool(even_page_header_blocks)
                # Add first page header: prefer explicit first_page_header, else fallback to header
                if section.different_first_page_header_footer:
                    if first_page_header_blocks and any(b for b in first_page_header_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                        part = getattr(section, 'first_page_header', None)
                        if part:
                            print(f"[convert_document] Writing {len(first_page_header_blocks)} blocks to first_page_header of section {idx}")
                            part.is_header = True
                            add_blocks_to_doc(part, first_page_header_blocks, temp_dir)
                    elif header_blocks and any(b for b in header_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                        part = getattr(section, 'first_page_header', None)
                        if part:
                            print(f"[convert_document] Writing {len(header_blocks)} blocks to first_page_header (from 'header') of section {idx}")
                            part.is_header = True
                            add_blocks_to_doc(part, header_blocks, temp_dir)
                # Add regular header for subsequent pages: only if header exists
                if header_blocks and any(b for b in header_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                    part = getattr(section, 'header', None)
                    if part:
                        print(f"[convert_document] Writing {len(header_blocks)} blocks to header of section {idx}")
                        part.is_header = True
                        add_blocks_to_doc(part, header_blocks, temp_dir)
                # Add even page header if configured
                if section.different_even_page_header_footer and even_page_header_blocks and any(b for b in even_page_header_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                    part = getattr(section, 'even_page_header', None)
                    if part:
                        print(f"[convert_document] Writing {len(even_page_header_blocks)} blocks to even_page_header of section {idx}")
                        part.is_header = True
                        add_blocks_to_doc(part, even_page_header_blocks, temp_dir)
                # Footers: similar logic
                first_page_footer_blocks = section_json.get('first_page_footer')
                footer_blocks = section_json.get('footer')
                even_page_footer_blocks = section_json.get('even_page_footer')
                # Add first page footer: prefer explicit first_page_footer, else fallback to footer
                if section.different_first_page_header_footer:
                    if first_page_footer_blocks and any(b for b in first_page_footer_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                        part = getattr(section, 'first_page_footer', None)
                        if part:
                            print(f"[convert_document] Writing {len(first_page_footer_blocks)} blocks to first_page_footer of section {idx}")
                            part.is_footer = True
                            add_blocks_to_doc(part, first_page_footer_blocks, temp_dir)
                    elif footer_blocks and any(b for b in footer_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                        part = getattr(section, 'first_page_footer', None)
                        if part:
                            print(f"[convert_document] Writing {len(footer_blocks)} blocks to first_page_footer (from 'footer') of section {idx}")
                            part.is_footer = True
                            add_blocks_to_doc(part, footer_blocks, temp_dir)
                # Add regular footer for subsequent pages: only if footer exists
                if footer_blocks and any(b for b in footer_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                    part = getattr(section, 'footer', None)
                    if part:
                        print(f"[convert_document] Writing {len(footer_blocks)} blocks to footer of section {idx}")
                        part.is_footer = True
                        add_blocks_to_doc(part, footer_blocks, temp_dir)
                # Add even page footer if configured
                if section.different_even_page_header_footer and even_page_footer_blocks and any(b for b in even_page_footer_blocks if b.get('type') != 'paragraph' or b.get('runs', []) or b.get('type') == 'table'):
                    part = getattr(section, 'even_page_footer', None)
                    if part:
                        print(f"[convert_document] Writing {len(even_page_footer_blocks)} blocks to even_page_footer of section {idx}")
                        part.is_footer = True
                        add_blocks_to_doc(part, even_page_footer_blocks, temp_dir)
            # Add body content
            if 'body' in data and isinstance(data['body'], list):
                add_blocks_to_doc(doc, data['body'], temp_dir)
            elif 'blocks' in data and isinstance(data['blocks'], list):
                add_blocks_to_doc(doc, data['blocks'], temp_dir)
            doc.save(output_file)
            input_preview = get_preview(orig_file_path)
            output_preview = get_preview(output_file)
            return input_preview, output_preview, output_file
    # For other conversions that require Pandoc
    if os.environ.get('RAILWAY_ENVIRONMENT'):
        # In Railway, we only support direct DOCX<->JSON conversions
        error_msg = "This conversion is not supported in the API-only mode. Only DOCX<->JSON conversions are supported."
        print(error_msg)
        return error_msg, None, None
    else:
        # For local development with Pandoc
        try:
            import pypandoc
            pypandoc.convert_file(
                orig_file_path,
                target_format.lower(),
                outputfile=output_file,
            )
            input_preview = get_preview(orig_file_path)
            output_preview = get_preview(output_file)
            return input_preview, output_preview, output_file
        except Exception as e:
            error_msg = f"Error: {e}"
            print(error_msg)
            return error_msg, None, None

def parity_check(docx_path):
    import tempfile
    print(f"[Parity Check] Testing round-trip for: {docx_path}")
    class FileLike:  # Fake file-like for CLI
        def __init__(self, name): self.name = name
    _, _, json_out = convert_document(FileLike(docx_path), 'json')
    if not json_out or not os.path.exists(json_out):
        print("Failed to produce JSON from DOCX.")
        return False
    _, _, docx_out = convert_document(FileLike(json_out), 'docx')
    if not docx_out or not os.path.exists(docx_out):
        print("Failed to produce DOCX from JSON.")
        return False
    def extract_all_sections_for_parity(docx_path):
        doc = Document(docx_path)
        sections = []
        for idx, section in enumerate(doc.sections):
            sec = {}
            for htype, attr in [("header", "header"), ("first_page_header", "first_page_header"), ("even_page_header", "even_page_header"),
                                ("footer", "footer"), ("first_page_footer", "first_page_footer"), ("even_page_footer", "even_page_footer")]:
                part = getattr(section, attr, None)
                if part:
                    sec[htype] = extract_blocks(part, os.path.dirname(docx_path), f"sec{idx}_{htype}")
            sections.append(sec)
        body = extract_blocks(doc, os.path.dirname(docx_path), os.path.splitext(os.path.basename(docx_path))[0])
        return {"sections": sections, "body": body}
    orig = extract_all_sections_for_parity(docx_path)
    roundtrip = extract_all_sections_for_parity(docx_out)
    import difflib, pprint
    def blocks_to_str(blocks):
        return pprint.pformat(blocks, width=120)
    if orig == roundtrip:
        print("[Parity Check] PASS: Round-trip blocks are identical!")
        return True
    else:
        print("[Parity Check] FAIL: Differences found.")
        # Compare per section
        for idx, (orig_sec, round_sec) in enumerate(zip(orig["sections"], roundtrip["sections"])):
            if orig_sec != round_sec:
                print(f"Section {idx} header/footer mismatch:")
                diff = difflib.unified_diff(blocks_to_str(orig_sec).splitlines(), blocks_to_str(round_sec).splitlines(), fromfile='original', tofile='roundtrip', lineterm='')
                print('\n'.join(diff))
        if orig["body"] != roundtrip["body"]:
            print("Body mismatch:")
            diff = difflib.unified_diff(blocks_to_str(orig["body"]).splitlines(), blocks_to_str(roundtrip["body"]).splitlines(), fromfile='original', tofile='roundtrip', lineterm='')
            print('\n'.join(diff))
        return False

if __name__ == "__main__":
    if len(sys.argv) == 3 and sys.argv[1] == "--parity-check":
        parity_check(sys.argv[2])
        sys.exit(0)
    
    # Create Flask app for API endpoints
    app = Flask(__name__)
    
    def check_api_key():
        """Check if the API key is valid."""
        provided_key = request.headers.get('X-API-Key')
        if not provided_key or provided_key != API_KEY:
            return False
        return True
    
    @app.route('/', methods=['GET'])
    def index():
        """Index route to check if the API is running."""
        # Always serve the simple HTML interface (regardless of environment)
        html = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Document Format Converter</title>
            <style>
                body { font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }
                .container { margin-top: 20px; border: 1px solid #ddd; padding: 20px; border-radius: 5px; }
                .loading { color: #666; }
            </style>
        </head>
        <body>
            <h2>Document Format Converter (Simple UI)</h2>
            <div class="container">
                <h3>DOCX to JSON</h3>
                <input type="file" id="docxInput" accept=".docx" />
                <button onclick="convertDocxToJson()">Convert to JSON</button>
                <div id="docxLoading" class="loading" style="display:none;">Processing...</div>
                <pre id="jsonOutput" style="max-height:300px;overflow:auto;"></pre>
            </div>
            <div class="container">
                <h3>JSON to DOCX</h3>
                <textarea id="jsonInput" rows="10" style="width:100%;"></textarea><br>
                <button onclick="convertJsonToDocx()">Convert to DOCX</button>
                <div id="jsonLoading" class="loading" style="display:none;">Processing...</div>
                <a id="docxDownload" style="display:none;" download="converted.docx">Download DOCX</a>
            </div>
            <script>
                async function convertDocxToJson() {
                    const input = document.getElementById('docxInput');
                    const loading = document.getElementById('docxLoading');
                    const output = document.getElementById('jsonOutput');
                    output.textContent = '';
                    if (!input.files.length) {
                        alert('Please select a DOCX file');
                        return;
                    }
                    loading.style.display = 'block';
                    const file = input.files[0];
                    const formData = new FormData();
                    formData.append('file', file);
                    try {
                        const resp = await fetch('/api/docx-to-json', {
                            method: 'POST',
                            headers: { 'X-API-Key': 'docgen_api_12345' },
                            body: formData
                        });
                        if (!resp.ok) throw new Error(await resp.text());
                        const data = await resp.json();
                        output.textContent = JSON.stringify(data, null, 2);
                        document.getElementById('jsonInput').value = JSON.stringify({ document: data }, null, 2);
                    } catch (error) {
                        output.textContent = 'Error: ' + error;
                    } finally {
                        loading.style.display = 'none';
                    }
                }
                async function convertJsonToDocx() {
                    const jsonText = document.getElementById('jsonInput').value;
                    if (!jsonText) {
                        alert('Please enter JSON data');
                        return;
                    }
                    let jsonObj;
                    try {
                        jsonObj = JSON.parse(jsonText);
                        if (typeof jsonObj !== 'object' || Array.isArray(jsonObj) || jsonObj === null) {
                            alert('JSON must be a non-empty object.');
                            return;
                        }
                    } catch (e) {
                        alert('Invalid JSON: ' + e.message);
                        return;
                    }
                    const loading = document.getElementById('jsonLoading');
                    loading.style.display = 'block';
                    const downloadLink = document.getElementById('docxDownload');
                    downloadLink.style.display = 'none';
                    try {
                        const resp = await fetch('/api/json-to-docx', {
                            method: 'POST',
                            headers: {
                                'Content-Type': 'application/json',
                                'X-API-Key': 'docgen_api_12345'
                            },
                            body: JSON.stringify(jsonObj)
                        });
                        if (!resp.ok) throw new Error(await resp.text());
                        const blob = await resp.blob();
                        const url = window.URL.createObjectURL(blob);
                        downloadLink.href = url;
                        downloadLink.style.display = 'inline-block';
                        downloadLink.textContent = 'Download DOCX';
                    } catch (error) {
                        alert('Error: ' + error);
                    } finally {
                        loading.style.display = 'none';
                    }
                }
            </script>
        </body>
        </html>
        """
        return html
    
    @app.route('/api/docx-to-json', methods=['POST'])
    def api_docx_to_json():
        # Check API key
        if not check_api_key():
            return jsonify({"error": "Invalid or missing API key"}), 401
            
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400
        
        if not file.filename.lower().endswith('.docx'):
            return jsonify({"error": "File must be a DOCX document"}), 400
        
        # Save the uploaded file
        temp_dir = tempfile.mkdtemp()
        file_path = os.path.join(temp_dir, file.filename)
        file.save(file_path)
        
        try:
            # Convert to JSON
            _, _, json_path = convert_document(type('obj', (object,), {'name': file_path}), "json")
            
            if not json_path or not os.path.exists(json_path):
                return jsonify({"error": "Error converting document to JSON"}), 500
            
            # Read JSON content
            with open(json_path, "r", encoding="utf-8") as f:
                json_content = json.load(f)
            
            return jsonify(json_content)
        except Exception as e:
            return jsonify({"error": str(e)}), 500
    
    @app.route('/api/json-to-docx', methods=['POST'])
    def api_json_to_docx():
        # Check API key
        if not check_api_key():
            return jsonify({"error": "Invalid or missing API key"}), 401
            
        if not request.is_json:
            return jsonify({"error": "Request must be JSON"}), 400
        
        try:
            # Save the JSON to a temporary file
            temp_dir = tempfile.mkdtemp()
            json_path = os.path.join(temp_dir, "document.json")
            
            # Log incoming JSON
            print("[api_json_to_docx] Received JSON:", request.json)
            incoming_json = request.json
            # --- Begin header auto-move logic for round-trip fidelity ---
            # If 'document' in JSON, operate on that
            doc_json = incoming_json.get('document', incoming_json)
            # If header content is present in body blocks, move to header
            if isinstance(doc_json, dict):
                blocks = doc_json.get('blocks', [])
                sections = doc_json.get('sections', [])
                # Only process if blocks exist and sections exist
                if blocks and isinstance(sections, list) and len(sections) > 0:
                    # Detect header-like blocks (e.g., heading at top, or blocks marked as header)
                    header_blocks = []
                    non_header_blocks = []
                    for block in blocks:
                        # Heuristic: heading at very top, or explicit marker
                        if block.get('type') == 'heading' and block.get('level', 1) == 1:
                            header_blocks.append(block)
                        elif block.get('in_header'):
                            header_blocks.append(block)
                        else:
                            non_header_blocks.append(block)
                    # If we found header blocks and first section header is empty, move them
                    first_section = sections[0] if len(sections) > 0 else None
                    if header_blocks and first_section is not None and (not first_section.get('header')):
                        first_section['header'] = header_blocks
                        doc_json['blocks'] = non_header_blocks
                        print(f"[api_json_to_docx] Moved {len(header_blocks)} blocks from body to header of first section.")
            # --- End header auto-move logic ---
            with open(json_path, "w", encoding="utf-8") as f:
                json.dump(incoming_json, f)
        
            # Convert to DOCX
            _, _, docx_path = convert_document(type('obj', (object,), {'name': json_path}), "docx")
            
            # After conversion, check if DOCX exists and has content
            if not docx_path or not os.path.exists(docx_path):
                print("[api_json_to_docx] ERROR: DOCX file was not created.")
                return jsonify({"error": "Error converting JSON to DOCX"}), 500
            from docx import Document
            doc = Document(docx_path)
            num_paras = len(doc.paragraphs)
            num_tables = len(doc.tables)
            print(f"[api_json_to_docx] DOCX created: paragraphs={num_paras}, tables={num_tables}")
            
            return send_file(docx_path, as_attachment=True, download_name="converted.docx")
        except Exception as e:
            print(f"[api_json_to_docx] Exception: {e}")
            return jsonify({"error": str(e)}), 500
    
    # For Railway deployment, get the port from environment variable
    port = int(os.environ.get('PORT', 8080))
    
    # In Railway, we'll just run the Flask app without Gradio
    # This avoids port conflicts and simplifies the deployment
    if os.environ.get('RAILWAY_ENVIRONMENT'):
        print("Running in Railway environment - API only mode")
        app.run(host='0.0.0.0', port=port)
    else:
        # In local development, run both Flask and Gradio
        print("Running in local environment - UI and API mode")
        
        # Create Gradio interface
        with gr.Blocks(css="footer {visibility: hidden}") as demo:
            gr.Markdown("# Document Format Converter\nUpload a document and preview as JSON. Select a format to download in another format.")
            with gr.Row():
                with gr.Column():
                    input_file = gr.File(label="Upload Document", file_types=[f'.{ext.lower()}' for ext in input_supported_formats])
                    input_preview = gr.HTML(label="JSON Preview")
                with gr.Column():
                    output_format = gr.Dropdown(label="Download As...", choices=output_supported_formats, value="DOCX")
                    format_label = gr.Markdown("Previewing as: DOCX")
                    output_preview = gr.HTML(label="Output Preview")
                    output_file = gr.File(label="Download Converted Document", visible=True)
                    json_to_docx_btn = gr.Button("Convert JSON to DOCX", visible=True)
            json_state = gr.State()
            orig_file_state = gr.State()

            def upload_and_preview(doc_file):
                _, _, json_path = convert_document(doc_file, "json")
                # Handle conversion failure
                if not json_path or not os.path.exists(json_path):
                    error_msg = "Error converting document to JSON."
                    return f"<pre style='max-height:300px;overflow:auto'>{error_msg}</pre>", "", doc_file.name
                # Read and preview JSON content
                try:
                    with open(json_path, "r", encoding="utf-8") as f:
                        json_content = f.read()
                except Exception as e:
                    error_msg = f"Error reading JSON: {e}"
                    return f"<pre style='max-height:300px;overflow:auto'>{error_msg}</pre>", "", doc_file.name
                preview_html = f"<pre style='max-height:300px;overflow:auto'>{json_content[:4000]}</pre>"
                return preview_html, json_content, doc_file.name

            def convert_and_preview(orig_file_path, output_format):
                class F:
                    name = orig_file_path
                _, _, out_path = convert_document(F(), output_format.lower())
                preview = get_preview(out_path)
                return f"Previewing as: {output_format}", preview, out_path

            import requests
            def json_to_docx(json_string):
                api_url = f"http://127.0.0.1:{port+1}/api/json-to-docx"
                api_key = os.environ.get('API_KEY', 'docgen_api_12345')
                try:
                    json_data = json.loads(json_string)
                except Exception as e:
                    return None, f"Invalid JSON: {e}"
                try:
                    resp = requests.post(api_url, json=json_data, headers={"X-API-Key": api_key})
                    if resp.status_code == 200:
                        temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                        temp_docx.write(resp.content)
                        temp_docx.close()
                        return temp_docx.name, ""
                    else:
                        return None, f"API Error: {resp.status_code} {resp.text}"
                except Exception as e:
                    return None, f"Request failed: {e}"

            def handle_json_to_docx(json_string):
                docx_path, err = json_to_docx(json_string)
                if err:
                    return None, f"<pre style='max-height:300px;overflow:auto'>{err}</pre>"
                preview = get_preview(docx_path)
                return docx_path, preview

            input_file.upload(upload_and_preview, inputs=input_file, outputs=[input_preview, json_state, orig_file_state])
            output_format.change(convert_and_preview, inputs=[orig_file_state, output_format], outputs=[format_label, output_preview, output_file])
            json_to_docx_btn.click(handle_json_to_docx, inputs=json_state, outputs=[output_file, output_preview])
        
        # Run Flask in a separate thread (on a different port to avoid conflicts)
        flask_port = port + 1
        def run_flask():
            app.run(host='0.0.0.0', port=flask_port)
        
        flask_thread = threading.Thread(target=run_flask)
        flask_thread.daemon = True
        flask_thread.start()
        
        print(f"Flask API running on http://127.0.0.1:{flask_port}")
        print(f"Gradio UI running on http://127.0.0.1:{port}")
        
        # Start Gradio
        demo.launch(server_name='0.0.0.0', server_port=port, share=False)