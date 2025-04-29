from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Add a title
title = doc.add_heading('Sample Document', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add some paragraphs
p = doc.add_paragraph('This is a sample document created for testing the Document Format Converter API. ')
p.add_run('This text is bold.').bold = True
p.add_run(' This text is italic.').italic = True
p.add_run(' This text is underlined.').underline = True

# Add a heading
doc.add_heading('Section 1', level=2)

# Add more paragraphs
doc.add_paragraph('This is a paragraph in section 1. It contains regular text.')

# Add another heading
doc.add_heading('Section 2', level=2)

# Add a paragraph with different formatting
p = doc.add_paragraph()
run = p.add_run('This paragraph has text with different font size.')
run.font.size = Pt(14)

# Save the document
doc.save('sample_document.docx')
print("Sample DOCX file created: sample_document.docx")
